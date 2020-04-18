# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt
from datetime import datetime
import pandas as pd
import numpy as np
import collections
from docx.enum.text import WD_ALIGN_PARAGRAPH
from spellchecker import SpellChecker



spell = SpellChecker()

### Read Excel
excel_file_name = 'testExcel.xlsx'
df = pd.read_excel(excel_file_name, sheet_names='ACUCOrganizationDonationVerific')
# complete_df = df[df['Entry_Status'] == 'Complete']
complete_df = df.fillna(0) 
# complete_df['机构名称'].fillna('',inplace = True) 
num_rows = len(complete_df.index)


### Date
now = datetime.now().strftime('%m_%d_%Y %Hh_%Mm_%Ss')
update_date = datetime.now().strftime('%m/%d/%Y %H:%M:%S')


### Write Word
document = Document()
# document.styles['Normal'].font.name = 'SimHei'
document.styles['Normal'].font.name = 'SimHei'

p = document.add_paragraph()
p_run = p.add_run('Org Name List Spell Checker')
p2= document.add_paragraph('Last Update: ' + str(update_date))
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_run.font.size = Pt(24)



table = document.add_table(rows=1, cols=1)
table.style = 'Table Grid'


# i = 0
item = ""
for index, row in complete_df.iterrows():
	first_row =  str(int(row['OrganizationSignUpListNumber您的机构在接龙里的序号'])) + '. ' + row['OrganizationNameInEnglish']
	words = spell.split_words(first_row)
	words_book = [spell.correction(word) for word in words]
	spell_check_result = spell.unknown(words_book)

	item = item + first_row + '\n' + str(spell_check_result) + '\n\n'

### breakdown

cell = table.cell(0,0)
cell.text = item


document.save('./output/Org Name List ' + now + '.docx')
print('Word file generate successful!')


