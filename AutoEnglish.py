# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt
from datetime import datetime
import pandas as pd
import numpy as np
import collections
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os





### Read Excel
excel_file_name = 'testExcel.xlsx'
df = pd.read_excel(excel_file_name, sheet_names='ACUCOrganizationDonationVerific')
# complete_df = df[df['Entry_Status'] == 'Complete']
complete_df = df.fillna(0) 
# complete_df['机构名称'].fillna('',inplace = True) 
num_rows = len(complete_df.index)


#Dict
translate_dict = {
'TotalPiecesOfN95InUS' : 'N95 Mask (Piece)',
'TotalPiecesOfSurgicalMaskInUS' : 'Medical Mask (Piece)',
'TotalPiecesOfMedicalProtectiveGownInUS' : 'Protective Garment',
'_3TotalPiecesOfTestKitsInUS' : 'Test Kits (set)',
'TotalPiecesOfMechanicalVentilator' : 'Ventilator',
'TotalPiecesOfHandSanitizerOrHandSoapInUS' : 'Hand Sanitizer (package)',
'TotalPiecesOfMedicalProtectiveHatInUS' : 'Protective Hat',
'TotalPiecesOfShoesCover' : 'Protective Shoes Cover',
'Gloves' : 'Protective Gloves',
'TotalPiecesOfGogglesInUS' : 'Goggles',
'TotalPiecesOfFaceShieldInUS' : 'Face Shield',
'DisinfectWipes' : 'Disinfect Wipes',
'TotalQtOfMeals' : 'Meals',
'_10TotalPiecesOfProtectiveKitsInUS' : 'ProtectiveKits',
# '_15OtherInKindSuppliesOrDonationsInUS' : 'Others'
}

summary_dict = {
'N95 Mask (Piece)' : 0,
'Medical Mask (Piece)' : 0,
'Protective Garment' : 0,
'Test Kits (set)' : 0,
'Ventilator' : 0,
'Hand Sanitizer (package)' : 0,
'Protective Hat' : 0,
'Protective Shoes Cover' : 0,
'Protective Gloves' : 0,
'Goggles' : 0,
'Face Shield' : 0,
'Disinfect Wipes' : 0,
'Meals' : 0,
'ProtectiveKits' : 0,
}

### Date
now = datetime.now().strftime('%m_%d_%Y %Hh_%Mm_%Ss')
update_date = datetime.now().strftime('%m/%d/%Y %H:%M:%S')


### Write Word
document = Document()
# document.styles['Normal'].font.name = 'SimHei'
document.styles['Normal'].font.name = 'SimHei'

p = document.add_paragraph()
p_run = p.add_run('ACUC Covid19 Donation Summary')
p2= document.add_paragraph('Last Update: ' + str(update_date))
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_run.font.size = Pt(24)



table = document.add_table(rows=num_rows + 2, cols=1)
table.style = 'Table Grid'

### Ranking and summary
rank_dict = {}
summary_amount = 0
summary_cash_amount = 0
summary_amount_not_arrived = 0  # 正在中国转运物资价值:
summary_amount_acuc_cash = 0 # 汇集ACUC现金捐赠
summary_amount_other_cash = 0 # 其余现金捐赠
summary_amount_supplies = 0 # 其他物资价值
summary_amount_outside_ny = 0 






for index, row in complete_df.iterrows():

	# first_row
	first_row = 'Registration number: ' + str(int(row['OrganizationSignUpListNumber您的机构在接龙里的序号'])) + '. \n' + 'Organization name: '+row['OrganizationNameInEnglish']
	
	# Second_row
	cash_amount = row['CashDonationAmountThroughACUC'] + row['AmountOfDonatedCash']
	if row['OrganizationSignUpListNumber您的机构在接龙里的序号'] == 75 or row['OrganizationSignUpListNumber您的机构在接龙里的序号'] == 62:
		summary_amount_outside_ny += cash_amount*0.6
	summary_amount_acuc_cash += row['CashDonationAmountThroughACUC']
	summary_amount_other_cash += row['AmountOfDonatedCash']
	summary_cash_amount = summary_cash_amount + cash_amount
	second_row = ''
	if cash_amount != 0:
		second_row = 'Cash donation: $' + str('{:,.2f}'.format(cash_amount)) + '\n'
	

	# Third_row
	row_part = row.iloc[12:58]
	items = ''
	for header , v in row_part.iteritems():
		if v != 0 and header in translate_dict:
			items = items + ' ' + translate_dict[header] + ': ' + str('{:,}'.format(int(v))) + ';'
			summary_dict[translate_dict[header]] += v
		if v != 0 and header == '_15OtherInKindSuppliesOrDonationsInUS' and v != '0':
			items = items + ' ' + str(v) + ';'
	third_row = ''
	if items != '':
		third_row = 'Supplies donation:' + items + '\n'
	

	# fourth_row
	supplie_value = row['ValuePriceOfTotalN95'] + \
	row['ValuePriceOfTotalSurgicalMask'] + row['ValuePriceOfTotalTestKits'] + \
	row['ValuePriceOfTotalGown'] + row['ValuePriceOfTotalVentilator'] + \
	row['ValuePriceOfTotalHandSoapSanitizer'] + row['ValuePriceOfTotalProtectiveHat'] + \
	row['ValuePriceOfTotalShoesCover'] + row['ValuePriceOfTotalFaceShield'] + \
	row['ValuePriceOfTotalProtectiveKits'] + row['ValuePriceOfTotalGloves'] + \
	row['ValuePriceOfTotalGoggles'] + row['ValuePriceOfTotalDisinfectWipes'] + \
	row['ValuePriceOfTotalMeals'] + row['TotalOfOthersValuePrice']
	### Special request
	if row['OrganizationSignUpListNumber您的机构在接龙里的序号'] == 75 or row['OrganizationSignUpListNumber您的机构在接龙里的序号'] == 62:
		summary_amount_outside_ny += supplie_value *0.6
	summary_amount_supplies += supplie_value
	total_amount = cash_amount + supplie_value + row['ValuePriceOfYourPurchaseThatAreNotYetArrived']
	fourth_row = 'Total donation value: $' + str('{:,.2f}'.format(total_amount))

	# fifth_row
	fifth_row = ''
	if row['ValuePriceOfYourPurchaseThatAreNotYetArrived'] != 0:
		if row['OrganizationSignUpListNumber您的机构在接龙里的序号'] == 75 or row['OrganizationSignUpListNumber您的机构在接龙里的序号'] == 62:
			summary_amount_outside_ny += row['ValuePriceOfYourPurchaseThatAreNotYetArrived']*0.6
		fifth_row += 'Value of supplies being transshipped overseas: $' + str('{:,.2f}'.format(row['ValuePriceOfYourPurchaseThatAreNotYetArrived'])) + '\n'
		summary_amount_not_arrived += row['ValuePriceOfYourPurchaseThatAreNotYetArrived']


	# Notes:
	last_row = ''
	if row['OrganizationSignUpListNumber您的机构在接龙里的序号'] == 75 or row['OrganizationSignUpListNumber您的机构在接龙里的序号'] == 62:
		last_row = 'Remarks: About 40% of its total donations donates to the New York tri-state area.\n'


	text_fill = first_row + '\n' + fourth_row + '\n' + second_row  + fifth_row + third_row + last_row + '\n'
	summary_amount = summary_amount + total_amount
	rank_dict[text_fill] = total_amount


### Sorting
# sorted_rank_dict = collections.OrderedDict(sorted(rank_dict.items(), key=lambda x: x[1], reverse=True))

### Write to docx
### Summary
## Normal summary
summary_cell = table.cell(0,0)
summary_row1 = 'Total number of registered donation organizations: ' + str(num_rows) + '\n'
summary_row2 = 'Total donated value: $' + str('{:,.2f}'.format(summary_amount)) + '\n'
summary_row3 = 'Total value of supplies donation being transshipped overseas: $' + str('{:,.2f}'.format(summary_amount_not_arrived)) + '\n'
summary_row4 = 'Total value of supplies donation reached the U.S.: $' + str('{:,.2f}'.format(summary_amount - summary_amount_not_arrived)) + '\n'
summary_row5 = 'Total value of donation to NY Tri-state area: $' + str('{:,.2f}'.format(summary_amount - summary_amount_outside_ny)) + '\n'
summary_row6 = 'Total value of ACUC cash donation: $' + str('{:,.2f}'.format(summary_amount_acuc_cash)) + '\n'
summary_row7 = 'Total value of other cash donation: $' + str('{:,.2f}'.format(summary_amount_other_cash)) + '\n'
summary_row8 = 'Total value of supplies donation: $' + str('{:,.2f}'.format(summary_amount_supplies)) + '\n'
s_row = ''
for s_name, s_amount in summary_dict.items():
	if s_amount != 0:
		s_row = s_row + '\t' + s_name + ': ' + str('{:,}'.format(int(s_amount))) + ';\n'
summary_row9 = 'Total number of supplies donation: ' + '\n' + s_row + '\n'

summary_cell.text = summary_row1 + summary_row2 + summary_row3 + summary_row4 + summary_row5 + summary_row6 + summary_row7 + summary_row8 + summary_row9

## Tree
summary_cell = table.cell(1,0)
summary_row1 = '+- ACUC Donation Summary\n\t|\n\t|--' + 'Total number of registered donation organizations:' + str(num_rows) + '\n\t|\n\t|--'
summary_row2 = 'Total donated value: $' + str('{:,.2f}'.format(summary_amount)) + '\n\t\t|\n\t\t|--'
summary_row3 = 'Total value of supplies donation being transshipped overseas: $' + str('{:,.2f}'.format(summary_amount_not_arrived)) + '\n' + '\t\t|\n\t\t|--'
summary_row4 = 'Total value of supplies donation reached the U.S.: $' + str('{:,.2f}'.format(summary_amount - summary_amount_not_arrived)) + '\n' +'\t\t\t|\n\t\t\t|--' 
summary_row5 = 'Total value of cash donation: $' + str('{:,.2f}'.format(summary_amount_acuc_cash+ summary_amount_other_cash)) + '\n' + '\t\t\t|\t|\n\t\t\t|\t|--'
summary_row6 = 'Total value of ACUC cash donation: $' + str('{:,.2f}'.format(summary_amount_acuc_cash)) + '\n' + '\t\t\t|\t|\n\t\t\t|\t|--'
summary_row7 = 'Total value of other cash donation: $' + str('{:,.2f}'.format(summary_amount_other_cash)) + '\n' + '\t\t\t|\n\t\t\t|--'
summary_row8 = 'Total value of supplies donation: $' + str('{:,.2f}'.format(summary_amount_supplies)) + '\n'
s_row = ''
for s_name, s_amount in summary_dict.items():
	if s_amount != 0:
		s_row = s_row + '\t\t\t\t|\n\t\t\t\t|--' + s_name + ': ' + str('{:,}'.format(int(s_amount))) + '\n'
summary_row9 = s_row + '\n'

summary_cell.text = summary_row1 + summary_row2 + summary_row3 + summary_row4 + summary_row5 + summary_row6 + summary_row7 + summary_row8 + summary_row9



### breakdown
i = 2
for text, price in rank_dict.items():
	cell = table.cell(i,0)
	cell.text = text
	i = i + 1




currentDay = datetime.now().day
currentMonth = datetime.now().month


# check output dir
path = './output/'+ str("{:02d}".format(currentMonth)) + "{:02d}".format(currentDay) + '/'
if not os.path.isdir(path):
	os.mkdir(path)

# save word file
document.save(path + 'English ACUC Donation Summary ' + now + '.docx')
print('Word file generate successful!')

# tree structure testing, normal, backup

### Tree
# summary_cell = table.cell(0,0)
# summary_row1 = '+- ACUC Donation Summary\n\t|\n\t|--' + '总计注册捐赠组织:' + str(num_rows) + '家\n\t|\n\t|--'
# summary_row2 = '总募集捐赠价值: $' + str('{:,.2f}'.format(summary_amount)) + '\n\t\t|\n\t\t|--'
# summary_row3 = '正在海外转运物资价值: $' + str('{:,.2f}'.format(summary_amount_not_arrived)) + '\n' + '\t\t|\n\t\t|--'
# summary_row4 = '已经到达美国捐赠价值: $' + str('{:,.2f}'.format(summary_amount - summary_amount_not_arrived)) + '\n' +'\t\t\t|\n\t\t\t|--' 
# summary_row5 = '现金捐赠总价值: $' + str('{:,.2f}'.format(summary_amount_acuc_cash+ summary_amount_other_cash)) + '\n' + '\t\t\t|\t|\n\t\t\t|\t|--'
# summary_row6 = '汇集ACUC现金捐赠: $' + str('{:,.2f}'.format(summary_amount_acuc_cash)) + '\n' + '\t\t\t|\t|\n\t\t\t|\t|--'
# summary_row7 = '其余现金捐赠: $' + str('{:,.2f}'.format(summary_amount_other_cash)) + '\n' + '\t\t\t|\n\t\t\t|--'
# summary_row8 = '物资捐赠总价值: $' + str('{:,.2f}'.format(summary_amount_supplies)) + '\n'
# s_row = ''
# for s_name, s_amount in summary_dict.items():
# 	s_row = s_row + '\t\t\t\t|\n\t\t\t\t|--' + s_name + ': ' + str('{:,}'.format(int(s_amount))) + '\n'
# summary_row9 = s_row + '\n'

## Normal summary
# summary_cell = table.cell(0,0)
# summary_row1 = '总计注册捐赠组织: ' + str(num_rows) + '家\n'
# summary_row2 = '总募集捐赠价值: $' + str('{:,.2f}'.format(summary_amount)) + '\n'
# summary_row3 = '正在海外转运物资价值: $' + str('{:,.2f}'.format(summary_amount_not_arrived)) + '\n'
# summary_row4 = '已经到达美国捐赠价值: $' + str('{:,.2f}'.format(summary_amount - summary_amount_not_arrived)) + '\n'
# summary_row5 = '捐赠给 NY Tri-state area 价值: $' + str('{:,.2f}'.format(summary_amount - summary_amount_outside_ny)) + '\n'
# summary_row6 = '汇集ACUC现金捐赠: $' + str('{:,.2f}'.format(summary_amount_acuc_cash)) + '\n'
# summary_row7 = '其余现金捐赠: $' + str('{:,.2f}'.format(summary_amount_other_cash)) + '\n'
# summary_row8 = '其他物资价值: $' + str('{:,.2f}'.format(summary_amount_supplies)) + '\n'
# s_row = ''
# for s_name, s_amount in summary_dict.items():
# 	s_row = s_row + ' ' + s_name + ': ' + str('{:,}'.format(int(s_amount))) + ';\n'
# summary_row9 = '其中物资捐赠总计: ' + '\n' + s_row + '\n'




# 
# 1. 正在转运列改为数字
# 2. 改127号，[海外捐款]



# Style
# 1. Page Margin to Normal
# 2. Chinese font SimHei, English font Calibri
# 3. table not across page
# 4. table font size 12
# 5. table layout
# 6. table autofit



# 0418 Update Notes:
# 1. change [Surgical Mask] to [Medical Mask]
# 2. ignore row['_15OtherInKindSuppliesOrDonationsInUS'] '0' item


# Update Notes:
# 1. 所有物资数量使用千分符(,)显示
# 2. 在明细中，增加[海外转运捐赠价值]
# 3. 隐藏明细中[现金捐赠，物资捐赠，海外转运捐赠]为空的，对应标题。

# Notes: 
# 1. 同意中英文两行
# 2. 总捐助价值下上面，明细在下
# 3. 英文就可以，不需要翻译中文



# index
# total_value_price = ['ValuePriceOfTotalN95',
# 'ValuePriceOfTotalSurgicalMask',
# 'ValuePriceOfTotalTestKits',
# 'ValuePriceOfTotalGown',
# 'ValuePriceOfTotalVentilator',
# 'ValuePriceOfTotalHandSoapSanitizer',
# 'ValuePriceOfTotalProtectiveHat',
# 'ValuePriceOfTotalShoesCover',
# 'ValuePriceOfTotalFaceShield',
# 'ValuePriceOfTotalProtectiveKits',
# 'ValuePriceOfTotalGloves',
# 'ValuePriceOfTotalGoggles',
# 'ValuePriceOfTotalDisinfectWipes',
# 'ValuePriceOfTotalMeals',
# 'TotalValuePrice',
# 'ValuePriceOfYourPurchaseThatAreNotYetArrived']


###Update Notes
# 1. re-arrange sorting algorithm and Num.
# 2. cover all company
# 3. summary header
### Style
# 1. font SimHei
# 2. Table not across page
# 3. table horizontal line