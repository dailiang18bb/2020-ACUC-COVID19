# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt
from datetime import datetime
import pandas as pd
import numpy as np
import collections
from docx.enum.text import WD_ALIGN_PARAGRAPH




### Read Excel
excel_file_name = 'testExcel.xlsx'
df = pd.read_excel(excel_file_name, sheet_names='ACUCOrganizationDonationVerific')
# complete_df = df[df['Entry_Status'] == 'Complete']
complete_df = df.fillna(0) 
# complete_df['机构名称'].fillna('',inplace = True) 
num_rows = len(complete_df.index)


#Dict
translate_dict = {
'TotalPiecesOfN95InUS' : 'N95 Mask (Piece)',
'TotalPiecesOfSurgicalMaskInUS' : 'Surgical Mask (Piece)',
'TotalPiecesOfMedicalProtectiveGownInUS' : 'Protective Garment',
'_3TotalPiecesOfTestKitsInUS' : 'Test Kits (set)',
'TotalPiecesOfMechanicalVentilator' : 'Ventilator',
'TotalPiecesOfHandSanitizerOrHandSoapInUS' : 'Hand Sanitizer (package)',
'TotalPiecesOfMedicalProtectiveHatInUS' : 'Protective Hat',
'TotalPiecesOfShoesCover' : 'Protective Shoes Cover',
'Gloves' : 'Protective Gloves',
'TotalPiecesOfGogglesInUS' : 'Goggles',
'TotalPiecesOfFaceShieldInUS' : 'Face Shield',
'DisinfectWipes' : 'Disinfect Wipes',
'TotalQtOfMeals' : 'Meals',
'_10TotalPiecesOfProtectiveKitsInUS' : 'ProtectiveKits',
# '_15OtherInKindSuppliesOrDonationsInUS' : 'Others'
}

error_check_dict = {
'TotalPiecesOfN95InUS' : 'ValuePriceOfTotalN95',
'TotalPiecesOfSurgicalMaskInUS' : 'ValuePriceOfTotalSurgicalMask',
'TotalPiecesOfMedicalProtectiveGownInUS' : 'ValuePriceOfTotalGown',
'_3TotalPiecesOfTestKitsInUS' : 'ValuePriceOfTotalTestKits',
'TotalPiecesOfMechanicalVentilator' : 'ValuePriceOfTotalVentilator',
'TotalPiecesOfHandSanitizerOrHandSoapInUS' : 'ValuePriceOfTotalHandSoapSanitizer',
'TotalPiecesOfMedicalProtectiveHatInUS' : 'ValuePriceOfTotalProtectiveHat',
'TotalPiecesOfShoesCover' : 'ValuePriceOfTotalShoesCover',
'Gloves' : 'ValuePriceOfTotalGloves',
'TotalPiecesOfGogglesInUS' : 'ValuePriceOfTotalGoggles',
'TotalPiecesOfFaceShieldInUS' : 'ValuePriceOfTotalFaceShield',
'DisinfectWipes' : 'ValuePriceOfTotalDisinfectWipes',
'TotalQtOfMeals' : 'ValuePriceOfTotalMeals',
'_10TotalPiecesOfProtectiveKitsInUS' : 'ValuePriceOfTotalProtectiveKits',
# '_15OtherInKindSuppliesOrDonationsInUS' : 'TotalValuePrice'
}

# price_check_dict = {
# 'ValuePriceOfTotalN95' : { 'min_price' : 2.45, 'max_price' :4.55 },
# 'ValuePriceOfTotalSurgicalMask' : { 'min_price' : 0.42, 'max_price' :0.78 },
# 'ValuePriceOfTotalGown' : { 'min_price' : 10.5, 'max_price' :19.5},
# 'ValuePriceOfTotalTestKits' : { 'min_price' : 7, 'max_price' :13},
# 'ValuePriceOfTotalVentilator' : { 'min_price' : 24500, 'max_price' :45500},
# # 'ValuePriceOfTotalHandSoapSanitizer': { min_price : , max_price :},
# 'ValuePriceOfTotalProtectiveHat' : { 'min_price' : 0.021, 'max_price' :0.039},
# 'ValuePriceOfTotalShoesCover' : { 'min_price' : 0.021, 'max_price' :0.039},
# 'ValuePriceOfTotalGloves' : { 'min_price' : 0.035, 'max_price' :0.065},
# 'ValuePriceOfTotalGoggles' : { 'min_price' : 2.1, 'max_price' :3.9},
# 'ValuePriceOfTotalFaceShield' : { 'min_price' : 2.1, 'max_price' :3.9},
# 'ValuePriceOfTotalDisinfectWipes' : { 'min_price' : 3.5, 'max_price' :6.5},
# 'ValuePriceOfTotalMeals' : { 'min_price' : 7, 'max_price' :13},
# # 'ValuePriceOfTotalProtectiveKits',
# }

# +-0.7
price_check_dict = {
'ValuePriceOfTotalN95' : { 'min_price' : 1.05, 'max_price' :5.95 },
'ValuePriceOfTotalSurgicalMask' : { 'min_price' : 0.18, 'max_price' :1.02 },
'ValuePriceOfTotalGown' : { 'min_price' : 4.5, 'max_price' :25.5},
'ValuePriceOfTotalTestKits' : { 'min_price' : 3, 'max_price' :17},
'ValuePriceOfTotalVentilator' : { 'min_price' : 10500, 'max_price' :59500},
# 'ValuePriceOfTotalHandSoapSanitizer': { min_price : , max_price :},
'ValuePriceOfTotalProtectiveHat' : { 'min_price' : 0.009, 'max_price' :0.051},
'ValuePriceOfTotalShoesCover' : { 'min_price' : 0.009, 'max_price' :0.051},
'ValuePriceOfTotalGloves' : { 'min_price' : 0.015, 'max_price' :0.085},
'ValuePriceOfTotalGoggles' : { 'min_price' : 0.9, 'max_price' :5.1},
'ValuePriceOfTotalFaceShield' : { 'min_price' : 0.9, 'max_price' :5.1},
'ValuePriceOfTotalDisinfectWipes' : { 'min_price' : 1.5, 'max_price' :8.5},
'ValuePriceOfTotalMeals' : { 'min_price' : 3, 'max_price' :17},
# 'ValuePriceOfTotalProtectiveKits',
}



summary_dict = {
'N95 Mask (Piece)' : 0,
'Surgical Mask (Piece)' : 0,
'Protective Garment' : 0,
'Test Kits (set)' : 0,
'Ventilator' : 0,
'Hand Sanitizer (package)' : 0,
'Protective Hat' : 0,
'Protective Shoes Cover' : 0,
'Protective Gloves' : 0,
'Goggles' : 0,
'Face Shield' : 0,
'Disinfect Wipes' : 0,
'Meals' : 0,
'ProtectiveKits' : 0,
}

### Date
now = datetime.now().strftime('%m_%d_%Y %Hh_%Mm_%Ss')
update_date = datetime.now().strftime('%m/%d/%Y %H:%M:%S')


### Write Word
document = Document()
# document.styles['Normal'].font.name = 'SimHei'
document.styles['Normal'].font.name = 'SimHei'

p = document.add_paragraph()
p_run = p.add_run('Price Check')
p2= document.add_paragraph('Last Update: ' + str(update_date))
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_run.font.size = Pt(24)



table = document.add_table(rows=num_rows + 1, cols=1)
table.style = 'Table Grid'

### Ranking and summary
rank_dict = {}
summary_amount = 0
summary_cash_amount = 0
error_dict = {}

# i = 0
for index, row in complete_df.iterrows():
	# cell = table.cell(i,0)

	flag = True
	# first_row

	if row['机构名称'] != 0:
		first_row = '报名序号: ' + str(int(row['OrganizationSignUpListNumber您的机构在接龙里的序号'])) + '. ' + row['OrganizationNameInEnglish'] + '_' + row['机构名称']
	else:
		first_row = '报名序号: ' + str(int(row['OrganizationSignUpListNumber您的机构在接龙里的序号'])) + '. ' + row['OrganizationNameInEnglish']
	# Second_row
	cash_amount = row['CashDonationAmountThroughACUC'] + row['AmountOfDonatedCash']
	summary_cash_amount = summary_cash_amount + cash_amount
	if cash_amount != 0:
		second_row = '现金捐款： $' + str('{:,.2f}'.format(cash_amount))
	else:
		second_row = '现金捐款： ' 

	# Third_row
	row_part = row.iloc[12:58]
	items = ''
	for header , v in row.iteritems():
		if v != 0 and header in translate_dict:
			if row[error_check_dict[header]] != 0:
				#error
				# items = items + ' \n' + '<<<ERROR Item found, missing value price of => [' + translate_dict[header] + ':' + str(int(v)) + '] >>>;\n'
				# flag = False
				if error_check_dict[header] in price_check_dict and (row[error_check_dict[header]] <= (price_check_dict[error_check_dict[header]]['min_price'] * v) or row[error_check_dict[header]] >= (price_check_dict[error_check_dict[header]]['max_price']) * v):
					#error
					items = items + ' \n' + '<<<ERROR Price Item found, [' + translate_dict[header] + ':' + str(int(v)) + '], filled price : $' + str('{:,.2f}'.format(row[error_check_dict[header]])) + ' >>>;\n'
					items = items + '==> Confidence unit price : ' + 'min_price: $' + str(price_check_dict[error_check_dict[header]]['min_price']) + ', max_price: $' + str(price_check_dict[error_check_dict[header]]['max_price']) + '\n'
					items = items + '==> Confidence price interval : ' + 'min: $' + str('{:,.2f}'.format(price_check_dict[error_check_dict[header]]['min_price'] * v)) + ', max: $' + str('{:,.2f}'.format(price_check_dict[error_check_dict[header]]['max_price'] * v)) + '\n'
					flag = False
				else:
					items = items + ' ' + translate_dict[header] + ':' + str(int(v)) + ';'



			summary_dict[translate_dict[header]] += v
			#add to row
		if v != 0 and header == '_15OtherInKindSuppliesOrDonationsInUS':
			items = items + ' ' + str(v) + ';'

			
	if flag == False:
		if row['TotalOfOthersValuePrice'] != 0:
			items = items + '\n[TotalOfOthersValuePrice] column filled: ' + str('{:,.2f}'.format(row['TotalOfOthersValuePrice']))
		else:
			items = items + '\n[TotalOfOthersValuePrice] column filled: EMPTY'

	third_row = '物品捐款：' + items

	# fourth_row
	total_amount = cash_amount + row['ValuePriceOfTotalN95'] + \
	row['ValuePriceOfTotalSurgicalMask'] + row['ValuePriceOfTotalTestKits'] + \
	row['ValuePriceOfTotalGown'] + row['ValuePriceOfTotalVentilator'] + \
	row['ValuePriceOfTotalHandSoapSanitizer'] + row['ValuePriceOfTotalProtectiveHat'] + \
	row['ValuePriceOfTotalShoesCover'] + row['ValuePriceOfTotalFaceShield'] + \
	row['ValuePriceOfTotalProtectiveKits'] + row['ValuePriceOfTotalGloves'] + \
	row['ValuePriceOfTotalGoggles'] + row['ValuePriceOfTotalDisinfectWipes'] + \
	row['ValuePriceOfTotalMeals'] + row['TotalOfOthersValuePrice'] 
	# row['ValuePriceOfYourPurchaseThatAreNotYetArrived']
	fourth_row = '总捐款价值： $' + str('{:,.2f}'.format(total_amount)) + '\n'
	text_fill = first_row + '\n' + second_row + '\n' + third_row + '\n' + fourth_row
	# cell.text = text_fill
	# i=i+1
	summary_amount = summary_amount + total_amount
	if flag == False:
		rank_dict[text_fill] = total_amount


### Sorting
# sorted_rank_dict = collections.OrderedDict(sorted(rank_dict.items(), key=lambda x: x[1], reverse=True))

### Write to docx
### Summary
summary_cell = table.cell(0,0)
summary_row1 = '总计注册捐赠组织： ' + str(num_rows) + '家\n'
summary_row2 = '总计捐赠价值： $' + str('{:,.2f}'.format(summary_amount)) + '\n'
summary_row3 = '其中现金捐赠价值： $' + str('{:,.2f}'.format(summary_cash_amount)) + '\n'
s_row = ''
for s_name, s_amount in summary_dict.items():
	s_row = s_row + ' ' + s_name + ':' + str(int(s_amount)) + ';'
summary_row4 = '其中物资捐赠总计： ' + s_row + '\n'
summary_cell.text = summary_row1 + summary_row2 + summary_row3 + summary_row4

### breakdown
i = 1
for text, price in rank_dict.items():
	cell = table.cell(i,0)
	cell.text = text
	i = i + 1




document.save('./output/Price Check ' + now + '.docx')
print('Word file generate successful!')






# index
# total_value_price = ['ValuePriceOfTotalN95',
# 'ValuePriceOfTotalSurgicalMask',
# 'ValuePriceOfTotalTestKits',
# 'ValuePriceOfTotalGown',
# 'ValuePriceOfTotalVentilator',
# 'ValuePriceOfTotalHandSoapSanitizer',
# 'ValuePriceOfTotalProtectiveHat',
# 'ValuePriceOfTotalShoesCover',
# 'ValuePriceOfTotalFaceShield',
# 'ValuePriceOfTotalProtectiveKits',
# 'ValuePriceOfTotalGloves',
# 'ValuePriceOfTotalGoggles',
# 'ValuePriceOfTotalDisinfectWipes',
# 'ValuePriceOfTotalMeals',
# 'TotalValuePrice',
# 'ValuePriceOfYourPurchaseThatAreNotYetArrived']


###Update Notes
# 1. re-arrange sorting algorithm and Num.
# 2. cover all company
# 3. summary header
### Style
# 1. font SimHei
# 2. Table not across page
# 3. table horizontal line