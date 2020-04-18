# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt
from datetime import datetime
import pandas as pd
import numpy as np
import collections
from docx.enum.text import WD_ALIGN_PARAGRAPH




### Read Excel
excel_file_name = 'testExcel.xlsx'
df = pd.read_excel(excel_file_name, sheet_names='ACUCOrganizationDonationVerific')
# complete_df = df[df['Entry_Status'] == 'Complete']
complete_df = df.fillna(0) 
# complete_df['机构名称'].fillna('',inplace = True) 
num_rows = len(complete_df.index)


### Date
now = datetime.now().strftime('%m_%d_%Y %Hh_%Mm_%Ss')
update_date = datetime.now().strftime('%m/%d/%Y %H:%M:%S')


### Write Word
document = Document()
# document.styles['Normal'].font.name = 'SimHei'
document.styles['Normal'].font.name = 'SimHei'

p = document.add_paragraph()
p_run = p.add_run('Check Missing and Dupicate number')
p2= document.add_paragraph('Last Update: ' + str(update_date))
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_run.font.size = Pt(24)



table = document.add_table(rows=1, cols=1)
table.style = 'Table Grid'

### Ranking and summary
rank_dict = {}
summary_amount = 0
summary_cash_amount = 0
error_dict = {}
missing_num_list = []
dupicate_num_list = []

def find_missing(lst):
	return [x for x in range(lst[0], lst[-1]+1) if x not in lst]

def checkIfDuplicates(listOfElems):
	setOfElems = set()
	temp_list = []
	for elem in listOfElems:
		if elem in setOfElems:
			temp_list.append(elem)
		else:
			setOfElems.add(elem)
	return temp_list



for index, row in complete_df.iterrows():
	missing_num_list.append(int(row['OrganizationSignUpListNumber您的机构在接龙里的序号']))


cell = table.cell(0,0)
temp_text = 'Missing number: ' + str(find_missing(missing_num_list)) + '\n'
temp_text = temp_text + 'Dupicate Number: ' + str(checkIfDuplicates(missing_num_list)) + '\n'
cell.text = temp_text




document.save('./output/Check missing and dupicate num ' + now + '.docx')
print('Word file generate successful!')



